from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
import pdfplumber
import tempfile
import os
from pydantic import BaseModel
from bs4 import BeautifulSoup
from xhtml2pdf import pisa
from docx import Document

router = APIRouter(prefix="/pdf", tags=["PDF Engine"])

@router.post("/import")
async def import_pdf(file: UploadFile = File(...)):
    """
    Parse an uploaded PDF file and return the extracted text and basic structure.
    This can be injected directly into the Yjs document state on the frontend.
    """
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
        
    text_content = []
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_path = temp_file.name

    try:
        with pdfplumber.open(temp_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
    except Exception as e:
        os.remove(temp_path)
        raise HTTPException(status_code=500, detail=f"Failed to parse PDF: {str(e)}")
        
    os.remove(temp_path)
    
    return {
        "filename": file.filename,
        "content": "\n\n".join(text_content),
        "pages": len(text_content)
    }

class ExportRequest(BaseModel):
    title: str = "Untitled Document"
    html_content: str

@router.post("/export")
async def export_pdf(req: ExportRequest, background_tasks: BackgroundTasks):
    """
    Convert the collaborative document (HTML representation) back to a beautifully formatted PDF.
    This version uses xhtml2pdf (pure Python) for compatibility and reliability.
    """
    output_path = tempfile.mktemp(suffix=".pdf")
    
    # Wrap in a premium print HTML template
    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @page {
                size: letter;
                margin: 1in;
            }
            body {{
                font-family: Helvetica, Arial, sans-serif;
                color: #1e293b;
                line-height: 1.6;
                font-size: 11pt;
            }}
            h1 {{ font-size: 24pt; margin-bottom: 12pt; color: #0f172a; border-bottom: 1px solid #e2e8f0; padding-bottom: 6pt; }}
            h2 {{ font-size: 18pt; margin-top: 18pt; margin-bottom: 8pt; color: #1e293b; }}
            h3 {{ font-size: 14pt; margin-top: 14pt; margin-bottom: 6pt; color: #334155; }}
            p {{ margin-bottom: 10pt; text-align: justify; }}
            blockquote {{
                margin: 0 0 12pt 0;
                padding-left: 12pt;
                border-left: 3pt solid #6366f1;
                color: #475569;
                font-style: italic;
                background-color: #f8fafc;
                padding-top: 4pt;
                padding-bottom: 4pt;
            }}
            pre {{
                background-color: #f1f5f9;
                padding: 10pt;
                font-family: Courier, monospace;
                font-size: 9pt;
                border-radius: 4pt;
                white-space: pre-wrap;
                margin-bottom: 12pt;
                border: 1px solid #e2e8f0;
            }}
            code {{
                background-color: #f1f5f9;
                font-family: Courier, monospace;
                font-size: 9.5pt;
                padding: 1pt 3pt;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin-top: 12pt;
                margin-bottom: 12pt;
            }}
            th, td {{
                border: 1px solid #cbd5e1;
                padding: 8pt;
                text-align: left;
                font-size: 10pt;
            }}
            th {{
                background-color: #f1f5f9;
                font-weight: bold;
                color: #0f172a;
            }}
            ul, ol {{
                margin-bottom: 10pt;
                padding-left: 20pt;
            }}
            li {{
                margin-bottom: 4pt;
            }}
        </style>
    </head>
    <body>
        <h1>{req.title}</h1>
        {req.html_content}
    </body>
    </html>
    """

    try:
        with open(output_path, "w+b") as result_file:
            pisa_status = pisa.CreatePDF(html_template, dest=result_file)
            if pisa_status.err:
                raise Exception(f"xhtml2pdf error code: {pisa_status.err}")
    except Exception as e:
        if os.path.exists(output_path):
            os.remove(output_path)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")
        
    background_tasks.add_task(os.remove, output_path)
    
    # Safe filename for download
    safe_filename = "".join(c for c in req.title if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_") or "Document"
    
    return FileResponse(
        path=output_path, 
        filename=f"{safe_filename}.pdf",
        media_type="application/pdf"
    )

@router.post("/export/docx")
async def export_docx(req: ExportRequest, background_tasks: BackgroundTasks):
    """
    Convert the HTML content into a genuine MS Word (.docx) document.
    """
    output_path = tempfile.mktemp(suffix=".docx")
    
    try:
        doc = Document()
        
        # Add primary title
        doc.add_heading(req.title, level=0)
        
        soup = BeautifulSoup(req.html_content, 'html.parser')
        
        def process_run_styles(child, paragraph, bold=False, italic=False, underline=False):
            from bs4 import NavigableString
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                return
                
            tag = child.name
            new_bold = bold or (tag in ['strong', 'b'])
            new_italic = italic or (tag in ['em', 'i'])
            new_underline = underline or (tag in ['u'])
            
            if tag == 'code':
                for sub_child in child.children:
                    if isinstance(sub_child, NavigableString):
                        run = paragraph.add_run(str(sub_child))
                        run.font.name = 'Courier New'
                        run.bold = new_bold
                        run.italic = new_italic
                        run.underline = new_underline
                    else:
                        process_run_styles(sub_child, paragraph, new_bold, new_italic, new_underline)
            else:
                for sub_child in child.children:
                    process_run_styles(sub_child, paragraph, new_bold, new_italic, new_underline)

        def process_element(element):
            from bs4 import NavigableString
            if isinstance(element, NavigableString):
                text = str(element).strip()
                if text:
                    p = doc.add_paragraph()
                    p.add_run(text)
                return

            tag_name = element.name
            
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag_name[1])
                p = doc.add_heading('', level=level)
                for child in element.children:
                    process_run_styles(child, p)
            elif tag_name == 'p':
                p = doc.add_paragraph()
                for child in element.children:
                    process_run_styles(child, p)
            elif tag_name == 'blockquote':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = 360000  # 0.25 inches
                for child in element.children:
                    process_run_styles(child, p, italic=True)
            elif tag_name == 'pre':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = 180000
                code_text = element.get_text()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
            elif tag_name in ['ul', 'ol']:
                list_style = 'List Bullet' if tag_name == 'ul' else 'List Number'
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(style=list_style)
                    for child in li.children:
                        process_run_styles(child, p)
            elif tag_name == 'table':
                rows = element.find_all('tr', recursive=False)
                if rows:
                    max_cols = 0
                    for r in rows:
                        cols = len(r.find_all(['td', 'th'], recursive=False))
                        max_cols = max(max_cols, cols)
                    
                    if max_cols > 0:
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = 'Table Grid'
                        
                        for r_idx, r in enumerate(rows):
                            cells = r.find_all(['td', 'th'], recursive=False)
                            for c_idx, cell in enumerate(cells):
                                docx_cell = table.cell(r_idx, min(c_idx, max_cols - 1))
                                cell_p = docx_cell.paragraphs[0]
                                for child in cell.children:
                                    process_run_styles(child, cell_p)
            else:
                for child in element.children:
                    process_element(child)

        for top_el in soup.contents:
            process_element(top_el)
            
        doc.save(output_path)
    except Exception as e:
        if os.path.exists(output_path):
            os.remove(output_path)
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")
        
    background_tasks.add_task(os.remove, output_path)
    
    # Safe filename for download
    safe_filename = "".join(c for c in req.title if c.isalnum() or c in (" ", "_", "-")).strip().replace(" ", "_") or "Document"
    
    return FileResponse(
        path=output_path, 
        filename=f"{safe_filename}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
